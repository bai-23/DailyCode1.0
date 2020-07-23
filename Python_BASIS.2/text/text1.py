#import docx
'''
f1 = docx.Document('C:\\Users\\86136\\Desktop\\office高级应用\\考题\\5周_长文档排版操作题√\\Word长文档排版.docx')
for i in f1.paragraphs:
    print(i.text)

import jieba
word = '今天天气真好！'
print(jieba.lcut(word))#   分词
print(jieba.lcut(word, cut_all=True))
print(jieba.lcut_for_search(word))#  精确分词
'''
#查找词频
from docx import Document
from jieba import lcut
f1 = Document('C:\\Users\\86136\\Desktop\\office高级应用\\考题\\5周_长文档排版操作题√\\Word长文档排版.docx')
doc = ''
for i in f1.paragraphs:
    doc = doc + i.text
word = lcut(doc)
print(word)

word_count = {}
for j in word:
    if len(j) == 1:
        continue
    else:
        word_count[j] = word_count.get(j, 0) + 1
print(word_count)
for k,v in sorted(word_count.items(), key=lambda y:y[1], reverse=True):
    print(k,':',v)
